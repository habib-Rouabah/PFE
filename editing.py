from docx import Document
import pandas as pd

# Création du DataFrame à partir des nouvelles données
data = {
    'Birch': {
        'precision': [0.33,0.43, '-'],
        'recall': [0.59, 0.20, '-'],
        'f1-score': [0.42, 0.27, '0.36'],
        'support': [2086, 3167, '5253']
    }
}

# Création du document Word
doc = Document()

# Parcourir les données pour chaque modèle
for model, model_data in data.items():
    # Création du DataFrame à partir des données
    df = pd.DataFrame(model_data)

    # Ajout du titre du modèle
    doc.add_heading(model.upper(), level=1)

    # Ajout du tableau au document Word
    table = doc.add_table(df.shape[0] + 1, df.shape[1] + 1)  # Augmenter les dimensions pour inclure les en-têtes

    # Ajout des en-têtes de colonne
    table.cell(0, 0).text = ''
    for j, col_name in enumerate(df.columns):
        table.cell(0, j + 1).text = col_name

    # Remplissage du tableau avec les données
    for i, (index, row) in enumerate(df.iterrows()):
        table.cell(i + 1, 0).text = str(index)
        for j, value in enumerate(row):
            table.cell(i + 1, j + 1).text = str(value)

    # Ajout d'une ligne vide entre les modèles
    doc.add_paragraph()

# Enregistrement du document Word
doc.save('scores.docx')
